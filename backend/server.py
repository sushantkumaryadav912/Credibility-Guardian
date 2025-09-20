import os
import json
import requests
import tempfile
import mimetypes
from flask import Flask, request, jsonify
from flask_cors import CORS
from bs4 import BeautifulSoup
import google.generativeai as genai
from dotenv import load_dotenv
import logging
from urllib.parse import urlparse
from werkzeug.utils import secure_filename

# Document processing imports
import pdfplumber
from docx import Document
from striprtf.striprtf import rtf_to_text

# Load environment variables from .env file
load_dotenv()

frontend_origin = os.getenv("FRONTEND_ORIGIN", "http://localhost:5173")

# Initialize Flask app
app = Flask(__name__)
CORS(app, origins=[frontend_origin])

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configure Gemini API
try:
    genai.configure(api_key=os.environ["GOOGLE_API_KEY"])
    logger.info("Gemini API configured successfully")
except KeyError:
    logger.error("GOOGLE_API_KEY environment variable not set")
    print("="*50)
    print("ERROR: GOOGLE_API_KEY environment variable not set.")
    print("Please add GOOGLE_API_KEY to your .env file.")
    print("="*50)

# File upload configuration
UPLOAD_FOLDER = tempfile.gettempdir()
MAX_CONTENT_LENGTH = 10 * 1024 * 1024  # 10MB
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Allowed file extensions and MIME types
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'rtf'}
ALLOWED_MIME_TYPES = {
    'application/pdf': 'pdf',
    'application/msword': 'doc', 
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'text/plain': 'txt',
    'application/rtf': 'rtf',
    'text/rtf': 'rtf'
}

def allowed_file(filename, content_type):
    """Check if the uploaded file is allowed"""
    if '.' not in filename:
        return False
    
    extension = filename.rsplit('.', 1)[1].lower()
    mime_extension = ALLOWED_MIME_TYPES.get(content_type)
    
    return extension in ALLOWED_EXTENSIONS or mime_extension in ALLOWED_EXTENSIONS

def extract_pdf_text(file_path):
    """Extract text from PDF using pdfplumber"""
    try:
        text_content = ""
        with pdfplumber.open(file_path) as pdf:
            logger.info(f"Processing PDF with {len(pdf.pages)} pages")
            
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                        logger.info(f"Extracted text from page {page_num + 1}")
                    else:
                        logger.warning(f"No text found on page {page_num + 1}")
                except Exception as e:
                    logger.error(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
        
        if not text_content.strip():
            raise ValueError("No text could be extracted from the PDF. The document might be image-based or corrupted.")
        
        return text_content.strip()
        
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

def extract_docx_text(file_path):
    """Extract text from DOCX using python-docx"""
    try:
        doc = Document(file_path)
        text_content = ""
        
        logger.info(f"Processing DOCX with {len(doc.paragraphs)} paragraphs")
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content += cell.text + " "
                text_content += "\n"
        
        if not text_content.strip():
            raise ValueError("No text could be extracted from the Word document.")
        
        return text_content.strip()
        
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise ValueError(f"Failed to extract text from Word document: {str(e)}")

def extract_doc_text(file_path):
    """Extract text from legacy DOC files (fallback method)"""
    try:
        # For legacy .doc files, we'll try to use python-docx
        # Note: python-docx primarily works with .docx files
        # For better .doc support, you might need python-docx2txt or antiword
        doc = Document(file_path)
        text_content = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
        
        if not text_content.strip():
            raise ValueError("No text could be extracted from the DOC file. Please convert to DOCX format.")
        
        return text_content.strip()
        
    except Exception as e:
        logger.error(f"DOC extraction error: {e}")
        raise ValueError(f"Failed to extract text from DOC file: {str(e)}. Please try converting to DOCX format.")

def extract_txt_text(file_path):
    """Extract text from TXT files"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            text_content = file.read()
        
        if not text_content.strip():
            raise ValueError("The text file appears to be empty.")
        
        return text_content.strip()
        
    except UnicodeDecodeError:
        # Try different encodings
        encodings = ['latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text_content = file.read()
                    logger.info(f"Successfully read TXT file with {encoding} encoding")
                    return text_content.strip()
            except:
                continue
        raise ValueError("Could not decode the text file. Please ensure it's a valid text file.")
        
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        raise ValueError(f"Failed to read text file: {str(e)}")

def extract_rtf_text(file_path):
    """Extract text from RTF files"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            rtf_content = file.read()
        
        text_content = rtf_to_text(rtf_content)
        
        if not text_content.strip():
            raise ValueError("No text could be extracted from the RTF file.")
        
        return text_content.strip()
        
    except Exception as e:
        logger.error(f"RTF extraction error: {e}")
        raise ValueError(f"Failed to extract text from RTF file: {str(e)}")

def extract_text_from_document(file_path, file_extension, original_filename):
    """Main function to extract text from various document types"""
    logger.info(f"Extracting text from {file_extension.upper()} file: {original_filename}")
    
    try:
        if file_extension == 'pdf':
            return extract_pdf_text(file_path)
        elif file_extension == 'docx':
            return extract_docx_text(file_path)
        elif file_extension == 'doc':
            return extract_doc_text(file_path)
        elif file_extension == 'txt':
            return extract_txt_text(file_path)
        elif file_extension == 'rtf':
            return extract_rtf_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    except Exception as e:
        logger.error(f"Text extraction failed for {original_filename}: {e}")
        raise

def fetch_article_text(url: str) -> str | None:
    """
    Fetches and extracts the main text content from a given news article URL.
    
    Args:
        url: The URL of the news article.
    
    Returns:
        The extracted article text as a string, or None if it fails.
    """
    logger.info(f"Fetching article from URL: {url}")
    try:
        # Validate URL format
        parsed_url = urlparse(url)
        if not parsed_url.scheme or not parsed_url.netloc:
            raise ValueError("Invalid URL format")
            
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove unwanted elements
        for script in soup(["script", "style", "nav", "footer", "aside", "header"]):
            script.decompose()

        # Try to find main content first
        main_content = soup.find('article') or soup.find('main') or soup.find(class_=['content', 'post-content', 'entry-content'])
        
        if main_content:
            paragraphs = main_content.find_all('p')
        else:
            paragraphs = soup.find_all('p')

        article_text = ' '.join([p.get_text().strip() for p in paragraphs if p.get_text().strip()])

        if len(article_text) < 200:
            logger.warning("Extracted text is very short. The page might have a paywall or complex structure.")
            return None

        return article_text

    except requests.RequestException as e:
        logger.error(f"Error fetching the URL: {e}")
        return None
    except Exception as e:
        logger.error(f"An error occurred during scraping: {e}")
        return None

def call_gemini_api(prompt: str) -> str:
    """Makes a live API call to the Google Gemini model."""
    logger.info("Making LIVE GEMINI API CALL")
    try:
        model = genai.GenerativeModel('gemini-1.5-flash-latest')
        generation_config = genai.types.GenerationConfig(
            response_mime_type="application/json"
        )
        response = model.generate_content(prompt, generation_config=generation_config)
        return response.text
    except Exception as e:
        logger.error(f"An error occurred during the API call: {e}")
        return json.dumps({"error": "Failed to get a response from the AI model.", "details": str(e)})

def analyze_text_for_misinformation(article_text: str) -> dict:
    """Analyzes a given text for signs of misinformation using a generative AI model."""
    prompt_template = f"""
    You are an expert misinformation and propaganda analyst. Your task is to analyze the following text for manipulative language, logical fallacies, emotional triggers, and signs of bias.

    Based on the text provided below, perform a detailed analysis and return your findings as a JSON object with the following exact structure:
    {{
      "credibility_score": <An integer score from 0 (completely untrustworthy) to 100 (highly credible)>,
      "summary_of_claims": "<A neutral, one-sentence summary of the main claims made in the text>",
      "analysis": {{
        "overall_assessment": "<A brief, overall assessment of the text's credibility and tone.>",
        "manipulative_techniques": [
          {{
            "technique": "<The name of the manipulative technique found (e.g., 'Emotional Appeal', 'Sensationalism & Hype', 'Weak Appeal to Authority', 'Logical Fallacy')>",
            "explanation": "<A brief explanation of how this technique is being used in the text.>",
            "flagged_quote": "<The exact quote from the text that demonstrates this technique.>"
          }}
        ]
      }}
    }}

    Analyze the following text:
    --- TEXT START ---
    {article_text}
    --- TEXT END ---
    """

    response_text = call_gemini_api(prompt_template)

    try:
        analysis_result = json.loads(response_text)
        return analysis_result
    except json.JSONDecodeError as e:
        logger.error(f"Error decoding JSON from AI response: {e}")
        logger.error(f"Raw response received: {response_text}")
        return {"error": "Failed to parse the analysis from the AI model.", "raw_response": response_text}

@app.route('/analyze', methods=['POST'])
def analyze():
    """Main analysis endpoint that handles URL, text, and document analysis"""
    try:
        # Check if it's a multipart form (document upload)
        if request.content_type and 'multipart/form-data' in request.content_type:
            return handle_document_analysis()
        else:
            return handle_url_text_analysis()

    except Exception as e:
        logger.error(f"Unexpected error in /analyze endpoint: {e}")
        return jsonify({
            'error': 'Internal server error occurred',
            'details': str(e)
        }), 500

def handle_document_analysis():
    """Handle document upload and analysis"""
    try:
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Validate file type and size
        if not allowed_file(file.filename, file.content_type):
            return jsonify({
                'error': f'File type not supported. Allowed types: {", ".join(ALLOWED_EXTENSIONS)}'
            }), 400
        
        # Get file extension
        filename = secure_filename(file.filename)
        file_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}')
        temp_file_path = temp_file.name
        
        try:
            # Save uploaded file
            file.save(temp_file_path)
            logger.info(f"File saved temporarily: {temp_file_path}")
            
            # Extract text from document
            extracted_text = extract_text_from_document(temp_file_path, file_extension, filename)
            
            # Validate extracted text length
            if len(extracted_text) < 50:
                return jsonify({
                    'error': 'Extracted text is too short for meaningful analysis. Please ensure the document contains readable text.'
                }), 400
            
            logger.info(f"Successfully extracted {len(extracted_text)} characters from document")
            
            # Analyze the extracted text
            credibility_report = analyze_text_for_misinformation(extracted_text)
            
            # Check if analysis was successful
            if "error" in credibility_report:
                return jsonify({
                    'error': 'Analysis failed',
                    'details': credibility_report
                }), 500
            
            # Add metadata to the response
            credibility_report['analysis_type'] = 'document'
            credibility_report['document_info'] = {
                'filename': filename,
                'file_type': file_extension,
                'text_length': len(extracted_text),
                'content_preview': extracted_text[:200] + '...' if len(extracted_text) > 200 else extracted_text
            }
            
            logger.info("Document analysis completed successfully")
            return jsonify(credibility_report)
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_file_path)
                logger.info("Temporary file cleaned up")
            except Exception as cleanup_error:
                logger.warning(f"Failed to clean up temporary file: {cleanup_error}")

    except Exception as e:
        logger.error(f"Document analysis error: {e}")
        return jsonify({
            'error': 'Failed to process document',
            'details': str(e)
        }), 500

def handle_url_text_analysis():
    """Handle URL and text analysis"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No JSON data provided'}), 400
            
        analysis_type = data.get('type', '').lower()
        content = data.get('data', '')
        
        if not analysis_type or not content:
            return jsonify({
                'error': 'Invalid request format. Expected: {"type": "url|text", "data": "content"}'
            }), 400

        # Process based on type
        if analysis_type == 'url':
            logger.info(f"Processing URL analysis for: {content}")
            article_content = fetch_article_text(content)
            
            if not article_content:
                return jsonify({
                    'error': 'Could not extract content from the provided URL. Please check the URL or try a different one.'
                }), 400
                
            text_to_analyze = article_content
            
        elif analysis_type == 'text':
            logger.info("Processing direct text analysis")
            text_to_analyze = content
            
            if len(text_to_analyze.strip()) < 50:
                return jsonify({
                    'error': 'Text content is too short for meaningful analysis. Please provide at least 50 characters.'
                }), 400
                
        else:
            return jsonify({
                'error': 'Invalid analysis type. Must be either "url" or "text"'
            }), 400

        # Perform the analysis
        logger.info("Starting credibility analysis")
        credibility_report = analyze_text_for_misinformation(text_to_analyze)
        
        # Check if analysis was successful
        if "error" in credibility_report:
            return jsonify({
                'error': 'Analysis failed',
                'details': credibility_report
            }), 500
        
        # Add metadata to the response
        credibility_report['analysis_type'] = analysis_type
        credibility_report['original_input'] = content if analysis_type == 'url' else content[:200] + '...' if len(content) > 200 else content
        
        logger.info("Analysis completed successfully")
        return jsonify(credibility_report)

    except Exception as e:
        logger.error(f"URL/Text analysis error: {e}")
        return jsonify({
            'error': 'Failed to process request',
            'details': str(e)
        }), 500

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    try:
        # Test if Gemini API is configured
        api_configured = "GOOGLE_API_KEY" in os.environ
        return jsonify({
            'status': 'healthy',
            'service': 'Credibility Analyzer API',
            'api_configured': api_configured,
            'supported_formats': list(ALLOWED_EXTENSIONS),
            'max_file_size': f"{MAX_CONTENT_LENGTH // (1024*1024)}MB"
        })
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'error': str(e)
        }), 500

@app.route('/', methods=['GET'])
def root():
    """Root endpoint with API information"""
    return jsonify({
        'service': 'Credibility Analyzer API',
        'version': '2.0.0',
        'description': 'Advanced AI-powered tool for detecting misinformation and analyzing content credibility',
        'endpoints': {
            '/analyze': 'POST - Analyze text, URL, or document for credibility',
            '/health': 'GET - Health check',
            '/': 'GET - API information'
        },
        'supported_formats': {
            'text': 'Direct text input',
            'url': 'Web articles and news content',
            'documents': list(ALLOWED_EXTENSIONS)
        },
        'usage': {
            'url_analysis': {
                'method': 'POST',
                'endpoint': '/analyze',
                'body': {
                    'type': 'url',
                    'data': 'https://example.com/article'
                }
            },
            'text_analysis': {
                'method': 'POST', 
                'endpoint': '/analyze',
                'body': {
                    'type': 'text',
                    'data': 'Your text content here...'
                }
            },
            'document_analysis': {
                'method': 'POST',
                'endpoint': '/analyze',
                'content_type': 'multipart/form-data',
                'body': 'Form data with file field'
            }
        }
    })

@app.errorhandler(413)
def too_large(error):
    return jsonify({
        'error': f'File too large. Maximum size is {MAX_CONTENT_LENGTH // (1024*1024)}MB'
    }), 413

@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Endpoint not found'}), 404

@app.errorhandler(405)
def method_not_allowed(error):
    return jsonify({'error': 'Method not allowed'}), 405

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port, debug=False)
